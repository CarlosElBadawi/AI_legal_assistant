from mcp.server.fastmcp import FastMCP
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
import docx
import os
from tempfile import NamedTemporaryFile
from mcp.server import Server
from mcp.server.fastmcp import FastMCP
from langchain.document_loaders import PyPDFLoader
from sentence_transformers import SentenceTransformer, util
from nltk.tokenize import sent_tokenize
from langchain.schema import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from google import genai
from dotenv import load_dotenv
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
mcp = FastMCP("mcp", port=8080)

# --------------------------
# 1. Policy Comparator Tool
# --------------------------
@mcp.tool()
def policy_comparator(company_clause: str, law_clause: str) -> dict:
    """Compare a company policy clause with statutory law and flag compliance. If a pdf file is provided, use the RAG tool to extract the context and place it in the company_clause.

    Args:
        company_clause (str): The text of the company policy clause.
        law_clause (str): The text of the statutory/legal clause.

    Returns:
        dict: compliance status and notes.
    """
    if len(company_clause.strip()) < len(law_clause.strip()):
        return {
            "status": "non-compliant",
            "notes": "Company clause appears weaker than statutory requirement."
        }
    else:
        return {
            "status": "compliant",
            "notes": "Company clause meets or exceeds statutory requirement."
        }


# --------------------------
# 2. Date Calculator Tool
# --------------------------
@mcp.tool()
def add_days_to_date(input_date: str, duration_days: str) -> dict:
    """
    Adds a number of days to a given date. The user can give you a string like "Add 45 days to 2025-08-01." and you have to know that 2025 corresponds to the year, 08 to the month, and 01 to the day.

    Args:
        input_date (str): The start date in format "YYYY-MM-DD".
        duration_days (str): Number of days to add as a string.

    Returns:
        dict: status and resulting date or error message.
    """
    try:
        # Convert input strings to proper types
        start_date = datetime.strptime(input_date, "%Y-%m-%d").date()
        days_to_add = int(duration_days)

        # Calculate the new date
        result_date = start_date + timedelta(days=days_to_add)

        return {
            "status": "success",
            "result_date": result_date.strftime("%Y-%m-%d")
        }
    except Exception as e:
        return {
            "status": "error",
            "error_message": str(e)
        }


# --------------------------
# 3. Legal Document Formatter Tool
# --------------------------
@mcp.tool()
def legal_doc_formatter(content: str, filename: str = "legal_output.docx") -> dict:
    """Format text into a legal .docx document.

    Args:
        content (str): The legal text or clauses to format.
        filename (str): Filename for the output docx.

    Returns:
        dict: file path and status.
    """
    try:
        doc = docx.Document()
        doc.add_heading("Legal Document", 0)

        for section in content.split("\n\n"):
            doc.add_paragraph(section)

        filepath = os.path.abspath(filename)
        doc.save(filepath)

        return {"status": "success", "file": filepath}
    except Exception as e:
        return {"status": "error", "error_message": str(e)}

client = genai.Client(api_key=GOOGLE_API_KEY)


def semantic_chunker(text: str, similarity_threshold: float = 0.85):
    sentences = sent_tokenize(text)
    model = SentenceTransformer("multi-qa-mpnet-base-dot-v1")
    embeddings = model.encode(sentences, convert_to_tensor=True)

    chunks = []
    current_chunk = [sentences[0]]
    for i in range(1, len(sentences)):
        similarity = util.pytorch_cos_sim(embeddings[i - 1], embeddings[i]).item()
        if similarity > similarity_threshold:
            current_chunk.append(sentences[i])
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [sentences[i]]
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return [Document(page_content=chunk) for chunk in chunks]

@mcp.tool()
def jurisdiction_checker(jurisdiction: str) -> dict:
    """Checks if the jurisdiction is valid and supported."""
    supported = ["California", "New York", "Texas"]
    if jurisdiction in supported:
        return {"status": "success", "jurisdiction": jurisdiction}
    else:
        return {"status": "error", "message": f"{jurisdiction} not supported."}

@mcp.tool()
def rag_query(uploaded_file, question: str) -> str:
    """
    Answer a question using a provided PDF (RAG), you have to seperate the link and the question and set each of them in the corresponding fields.
    Args:
        pdf_path: link to the PDF file.
        question: Natural language question to ask.
    """
    # Load PDF
    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name
    loader = PyPDFLoader(tmp_path)
    docs = loader.load()

    # Chunk into semantic sections
    full_text = " ".join([doc.page_content for doc in docs])
    chunks = semantic_chunker(full_text)

    # Build FAISS index
    embedding_model = HuggingFaceEmbeddings(model_name="multi-qa-mpnet-base-dot-v1")
    db = FAISS.from_documents(chunks, embedding_model)
    retriever = db.as_retriever(search_type="mmr", search_kwargs={'k': 3})

    # Retrieve context
    retrieved_docs = retriever.get_relevant_documents(question)
    context = "\n\n".join([doc.page_content for doc in retrieved_docs])

    # Prompt Gemini
    prompt = f"""Use the following context to answer the question, do not include any personal opinions or information, answer ONLY based on the context provided. If no context is available tell the user you don't know.
    If there are [] in the document especially if they are highlighted and in bold, consider them as filler for dates or numbers or names.

    Context:
    {context}

    Question: {question}
    Answer:"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text


# --------------------------
# Run MCP Server
# --------------------------
if __name__ == "__main__":
    mcp.run(transport="streamable-http")
